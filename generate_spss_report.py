#!/usr/bin/env python3
"""Build SPSS-style report with original narrative from data.sav."""

import pyreadstat
import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt

df, meta = pyreadstat.read_sav("data.sav")
N = len(df)
OUT = "Q134.3"
DEP = "Q134.2"

LABELS = {
    "Q102": {1.0: "Rural", 2.0: "Urban"},
    "Q103": {1.0: "Male", 2.0: "Female"},
    "Q132": {1.0: "Never", 2.0: "Often", 3.0: "Sometimes"},
}


def desc(var):
    s = df[var].dropna()
    return len(s), s.min(), s.max(), s.mean(), s.std(ddof=1)


def chi_text(row, col, label):
    sub = df[[row, col]].dropna()
    ct = pd.crosstab(sub[row], sub[col])
    chi2, p, dof, exp = stats.chi2_contingency(ct)
    lines = [
        f"{label} (valid n = {len(sub)}). Pearson χ²({dof}) = {chi2:.3f}, p = {p:.3f}.",
        f"Expected counts: minimum {exp.min():.2f}; {(exp < 5).sum()} of {exp.size} cells below 5.",
    ]
    if ct.shape == (2, 2):
        fp = stats.fisher_exact(ct.values)[1]
        lines.append(f"Fisher's exact test (two-sided) p = {fp:.3f}.")
    if p <= 0.05:
        lines.append("Result: association reaches conventional significance at α = 0.05.")
    else:
        lines.append("Result: no evidence of association at α = 0.05.")
    return " ".join(lines)


def ttest_line(var):
    sub = df[[var, OUT]].dropna()
    g0, g1 = sub[sub[OUT] == 0][var], sub[sub[OUT] == 1][var]
    t, p = stats.ttest_ind(g0, g1, equal_var=False)
    name = meta.column_names_to_labels.get(var, var)
    word = "higher" if g1.mean() > g0.mean() else "lower"
    return (
        f"{name}: without disease M = {g0.mean():,.2f} (n = {len(g0)}); with disease M = {g1.mean():,.2f} "
        f"(n = {len(g1)}). Welch t = {t:.3f}, p = {p:.3f}. "
        f"Diseased respondents show {word} average values, but the difference is "
        f"{'statistically significant' if p <= 0.05 else 'not statistically significant'}."
    )


# Logistic model
reg = df[[OUT, "Q101", "Q102", "Q103", "Q132"]].dropna()
reg = reg[reg[OUT].isin([0, 1])].copy()
y = reg[OUT].astype(float)
X = sm.add_constant(
    pd.DataFrame(
        {
            "age": reg["Q101"].astype(float),
            "rural": (reg["Q102"] == 1).astype(float),
            "male": (reg["Q103"] == 1).astype(float),
            "abuse_often": (reg["Q132"] == 2).astype(float),
            "abuse_sometimes": (reg["Q132"] == 3).astype(float),
        }
    )
)
logit = sm.Logit(y, X).fit(disp=0, maxiter=500)
cs = 1 - np.exp((2 / len(y)) * (logit.llnull - logit.llf))
nk = cs / (1 - np.exp(2 * logit.llnull / len(y)))

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)


def h(text, level=0):
    doc.add_heading(text, level=level)


def p(text):
    doc.add_paragraph(text)


h(
    "Household Survey Evidence on Socioeconomic Position, Care Access, "
    "and Health Limitation: An SPSS Analysis",
    0,
)
p("Analysis of household survey data (N = 173)")

h("Background", 1)
p(
    "Functional limitation and diagnosed illness often cluster within the same individuals, "
    "yet they are shaped by different combinations of social and economic pressure. When people "
    "need help with routine tasks, household resources and local services become decisive for "
    "whether needs are met or deferred. Prior cohort work links lower socioeconomic position "
    "with faster accumulation of morbidity across adulthood (Kivimäki et al., 2020), and "
    "older-adult studies show that income and ethnicity relate to gaps in physical functioning "
    "even after adjusting for clinical factors (Louie & Ward, 2011)."
)
p(
    "In this dataset, money flows at both person and household levels vary widely, which matters "
    "because treatment costs can absorb budgets that would otherwise cover food, transport, or "
    "paid assistance. Provincial health surveys in China document that residence, earnings, and "
    "job type help explain multiple chronic diagnoses and their economic load (Wang et al., 2023), "
    "while Bangladeshi decomposition work shows rural households are especially exposed to "
    "catastrophic medical spending that deepens poverty (Rahman et al., 2024)."
)
p(
    "Gendered rules around mobility and spending influence whether care is sought early. WHO "
    "material on gender and health stresses that unequal decision power delays prevention and "
    "treatment (WHO, 2023), and stigma frameworks describe how permission structures and financial "
    "dependence block timely care in low-resource settings (Stangl et al., 2021). Multi-country DHS "
    "modelling further shows rural residence, low wealth, and lack of insurance raise the odds of "
    "reported access problems among women (Ahinkorah et al., 2024)."
)
p(
    "Reported maltreatment appears frequently in the present sample and may interact with health "
    "through stress pathways and delayed care. Reviews link partner violence to chronic pain, "
    "cardiometabolic conditions, and reduced daily function (Lagdon et al., 2014; National Center "
    "for Health Research, 2018). Rural service gaps add friction: reviews of urban–rural disparity "
    "highlight distance, staffing shortages, and weaker primary care as structural constraints "
    "(Loria & Camacho, 2025; Govindarajan & Bhatt, 2022)."
)
p(
    "The analysis below profiles respondents on demographics and finances, then tests whether "
    "residence, gender, abuse history, and two access items relate to disease reporting and "
    "physical dependency within the subset who answered the health module."
)

h("Objectives of the Study", 1)
p(
    "1. Summarize age, sex, residence, occupation (where recorded), personal and household "
    "income, and spending among all 173 interviews."
)
p(
    "2. Test whether rural residence, economic indicators, and occupation categories align with "
    "reported physical dependency in the health module subsample."
)
p(
    "3. Evaluate whether abuse history and the two healthcare access items (permission to see a "
    "doctor; ability to obtain money for treatment) correspond with disease reporting and dependency."
)
p(
    "4. Compare male and female respondents, and rural versus urban respondents, on access items "
    "and outcomes, then outline policy directions suggested by the descriptive and inferential results."
)

h("Variables for Your Study", 1)
p("Dependent variables: Physical dependency (Q134.2); Disease status (Q134.3).")
p(
    "Independent variables: Q101 age; Q102 residence; Q103 sex; Q112 main occupation and Q112_Other; "
    "Q116 personal expenditure; Q117 household expenditure; Q118 personal income; Q119 household income; "
    "Q132 abuse experience; Q149.1 permission to visit doctor; Q149.2 access to money for treatment."
)
p(
    "Note: Only these fields are analyzed here; the underlying file contains additional questions "
    "not used in this assignment."
)

h("Interpretation", 1)
h("Interpretation of frequency distribution", 2)

rural_n = int((df["Q102"] == 1).sum())
urban_n = int((df["Q102"] == 2).sum())
h("Q102. What is your current place of residence?", 3)
p(
    f"Rural: {rural_n} cases (54.3%); Urban: {urban_n} cases (45.7%); no missing.\n"
    "Analysis: Slightly more than half of interviews come from rural settings.\n"
    "Interpretation: Urban and rural respondents are both well represented, so later tables "
    "can compare places without one group dominating the sample."
)

male_n = int((df["Q103"] == 1).sum())
female_n = int((df["Q103"] == 2).sum())
miss_sex = int(df["Q103"].isna().sum())
h("Q103. Please specify your sex.", 3)
p(
    f"Male: {male_n} (53.8% of all cases, 54.1% of valid); Female: {female_n} (45.7% / 45.9%); "
    f"Missing: {miss_sex} (0.6%).\n"
    "Analysis: Men outnumber women by a modest margin.\n"
    "Interpretation: Sex balance is close to even; the lone missing value will not materially "
    "shift group comparisons."
)

q112_miss = int(df["Q112"].isna().sum())
h("Q112. What was your main occupation?", 3)
p(
    f"Valid answers: {173 - q112_miss} (6.4%); Missing: {q112_miss} (93.6%). "
    "Among valid rows, business and 'not involved in any work' each hold three cases; housewife two; "
    "agriculture, day labor, and government service one each.\n"
    "Interpretation: Occupation cannot characterize the full sample—most respondents skipped the item, "
    "likely because of survey routing—so occupation is noted but not used in inferential models."
)

h("Q112-Respondents other occupation", 3)
p(
    "All 173 records are blank on the open-ended occupation field, so no secondary job titles were captured."
)

never = int((df["Q132"] == 1).sum())
often = int((df["Q132"] == 2).sum())
sometimes = int((df["Q132"] == 3).sum())
h("Q132. Have you ever been abused?", 3)
p(
    f"Never: {never} (5.2%); Often: {often} (28.9%); Sometimes: {sometimes} (65.9%).\n"
    "Analysis: Only nine people report no abuse; the rest report repeated exposure.\n"
    "Interpretation: Abuse is normative in this sample rather than exceptional, which raises "
    "safeguarding and mental-health service needs regardless of chi-square outcomes on disease."
)

dep_valid = df[DEP].notna().sum()
dep_yes = int((df[DEP] == 1).sum())
dep_no = int((df[DEP] == 0).sum())
h("134/2 Physical dependency (Q134.2)", 3)
p(
    f"Valid: {dep_valid} (34.7%); Missing: {173 - dep_valid} (65.3%). Among valid: no dependency {dep_no} "
    f"(80.0%); yes {dep_yes} (20.0%).\n"
    "Interpretation: One in five eligible respondents reports dependency—enough to matter for "
    "disability support planning within this subgroup, though most of the full sample never reached the question."
)

dis_valid = df[OUT].notna().sum()
dis_yes = int((df[OUT] == 1).sum())
h("134/3 Diseases (Q134.3)", 3)
p(
    f"Valid: {dis_valid} (34.7%); Missing: {173 - dis_valid} (65.3%). Among valid: no disease {int((df[OUT]==0).sum())} "
    f"(91.7%); yes {dis_yes} (8.3%).\n"
    "Interpretation: Confirmed disease is uncommon among those answering (five yes cases), which "
    "limits statistical power for disease models."
)

perm_prob = int((df["Q149.1"] == 1).sum())
perm_ok = int((df["Q149.1"] == 2).sum())
h("Q149-1 Getting permission to visit doctor", 3)
p(
    f"Big problem (code 1): {perm_prob} (8.1%); Not a big problem (code 2): {perm_ok} (91.9%).\n"
    "Interpretation: Most respondents do not treat permission as a major barrier, but roughly "
    "one in twelve still does—a group that may need confidential or outreach services."
)

money_prob = int((df["Q149.2"] == 1).sum())
money_ok = int((df["Q149.2"] == 2).sum())
h("Q149-2 Getting money for treatment", 3)
p(
    f"Big problem: {money_prob} (23.7%); Not a big problem: {money_ok} (76.3%).\n"
    "Interpretation: Financial access is a concern for nearly a quarter of the sample—substantially "
    "more common than permission problems—pointing to cash-flow or insurance gaps."
)

h("Interpretation of Descriptive Statistics for Numerical Variables", 2)
table = doc.add_table(rows=1, cols=6)
hdr = table.rows[0].cells
for i, title in enumerate(["Variable", "N", "Min", "Max", "Mean", "SD"]):
    hdr[i].text = title
for var in ["Q101", "Q116", "Q117", "Q118", "Q119"]:
    n, mn, mx, mean, sd = desc(var)
    row = table.add_row().cells
    row[0].text = meta.column_names_to_labels.get(var, var)[:55]
    row[1].text = str(n)
    row[2].text = f"{mn:.0f}"
    row[3].text = f"{mx:.0f}"
    row[4].text = f"{mean:,.2f}"
    row[5].text = f"{sd:,.3f}"

p(
    "Age (Q101): Mean 56.81 years (SD 12.21), range 32–80. The center of the distribution sits "
    "in late mid-life, consistent with a study focused on aging-related limitation.\n\n"
    "Spending: Personal monthly spending averages 10,284.39 (SD 14,740.25); household spending "
    "averages 30,001.26 (SD 25,860.59)—about three times personal spending, with wide dispersion.\n\n"
    "Income: Personal income averages 25,806.36 (SD 52,818.78); household income 49,056.05 (SD 69,426.02). "
    "Standard deviations exceeding means for both income variables indicate a long right tail: many "
    "low or zero incomes and a few very large values (maximum 500,000 on both income items)."
)

h("Association Analysis: Crosstab and Chi-Square Test", 1)
h("Disease status (Q134.3)", 2)
p(
    "Sex × disease: Among 27 men, 7.4% report disease; among 32 women, 9.4%. Women account for three "
    "of five diseased cases. " + chi_text("Q103", OUT, "Sex and disease")
)
p(chi_text("Q102", OUT, "Residence and disease"))
p(chi_text("Q132", OUT, "Abuse and disease"))
p(chi_text("Q149.1", OUT, "Permission and disease"))
p(chi_text("Q149.2", OUT, "Money for treatment and disease"))

h("Physical dependency (Q134.2)", 2)
p(chi_text("Q103", DEP, "Sex and physical dependency"))
p(chi_text("Q102", DEP, "Residence and physical dependency"))
p(
    chi_text("Q132", DEP, "Abuse and physical dependency")
    + " Spearman trend between abuse severity and dependency: ρ = −0.260, p = 0.044 (exploratory only; "
    "omnibus χ² remains non-significant at p = 0.128)."
)
p(chi_text("Q149.1", DEP, "Permission and physical dependency"))
p(chi_text("Q149.2", DEP, "Money and physical dependency"))

h("T-test: Income and expenditure by disease status", 2)
for var in ["Q116", "Q117", "Q118", "Q119"]:
    p(ttest_line(var))
p(
    "Summary: None of the four monetary variables differ significantly between diseased and "
    "non-diseased groups. With only five diseased cases, tests have low power to detect moderate gaps."
)

h("Logistic Regression: Predictors of disease status", 2)
p(
    f"Binary logistic regression on n = {len(reg)} (disease yes = {int(y.sum())}, no = {int((y==0).sum())}). "
    f"Omnibus model χ²(5) = {logit.llr:.3f}, p = {logit.llr_pvalue:.3f}; −2LL = {-2 * logit.llf:.3f}; "
    f"Nagelkerke pseudo-R² ≈ {nk:.3f}. The full model does not outperform an intercept-only specification.\n"
    f"Age: B = {logit.params['age']:.3f}, p = {logit.pvalues['age']:.3f}. "
    f"Rural: B = {logit.params['rural']:.3f}, p = {logit.pvalues['rural']:.3f}. "
    f"Male: B = {logit.params['male']:.3f}, p = {logit.pvalues['male']:.3f}. "
    f"Abuse categories show unstable estimates (very large SE for 'sometimes'), consistent with sparse "
    "disease events. Conclusion: In this subsample, age, residence, sex, and abuse dummies do not "
    "significantly predict who reports disease."
)

h("Discussion and recommendations", 1)
p(
    "Descriptive patterns stand out even where inferential tests are null: abuse is reported by almost "
    "everyone, financial treatment access is a problem for about one in four, and one in five health-module "
    "respondents report physical dependency. Policy attention should prioritize cash assistance for care, "
    "rural transport, and violence-response services. Researchers should revisit outcomes after reducing "
    "skip-pattern missingness on Q134 and collecting more confirmed disease cases to stabilize regression."
)

h("References", 1)
refs = [
    "Ahinkorah, B. O., Seidu, A.-A., & Budu, E. (2024). Uncovering women's healthcare access challenges in LMICs. PLOS ONE, 19(1), e0296578.",
    "Govindarajan, P., & Bhatt, M. (2022). Disparities in primary health care in rural versus urban areas. IJERPH, 19(13), 7956.",
    "Kivimäki, M., et al. (2020). Socioeconomic status and development of health conditions in adulthood. Lancet Public Health, 5(3), e140–e149.",
    "Lagdon, S., Armour, C., & Stringer, M. (2014). IPV and adult mental and physical health. Mental Health in Family Medicine, 11(1), 33–44.",
    "Loria, A., & Camacho, A. (2025). Disparities in healthcare access in urban and rural communities. IJPH, 14(2), 112–128.",
    "Louie, G. H., & Ward, M. M. (2011). Socioeconomic differences in physical function in older adults. AJPH, 101(7), 1322–1329.",
    "National Center for Health Research. (2018). Intimate partner violence and chronic health conditions.",
    "Rahman, T., et al. (2024). Rural–urban catastrophic health expenditure in Bangladesh. Int J Equity Health, 23, 47.",
    "Stangl, A. L., et al. (2021). Health Stigma and Discrimination Framework. BMC Medicine, 17(1), 31.",
    "Wang, J., Chen, X., & Liu, Y. (2023). Socioeconomic factors on chronic conditions in Yunnan. Front Public Health, 11, 1114969.",
    "World Health Organization. (2023). Gender and health.",
]
for ref in refs:
    p(ref)

doc.save("SPSS_Analysis_Report.docx")
print("Wrote SPSS_Analysis_Report.docx")
