import pyreadstat
import pandas as pd
from scipy import stats
import statsmodels.api as sm
from docx import Document

# LOAD DATA
df, meta = pyreadstat.read_sav("data.sav")

# =========================
# PICK CLEAN VARIABLES
# =========================
vars = ["Socialcapitalscore", "Meanscoresocialcapital", "Q129ADLs_score"]
df = df[vars]

# =========================
# DESCRIPTIVE STATS
# =========================
desc = df.describe()

# =========================
# CORRELATION MATRIX
# =========================
corr = df.corr()

# =========================
# REGRESSION
# Social capital predicts ADL score
# =========================
X = sm.add_constant(df["Socialcapitalscore"])
y = df["Q129ADLs_score"]

model = sm.OLS(y, X).fit()

# =========================
# WORD REPORT
# =========================
doc = Document()
doc.add_heading("SPSS STYLE REPORT", 0)

doc.add_heading("Descriptive Statistics", 1)
doc.add_paragraph(str(desc))

doc.add_heading("Correlation Matrix", 1)
doc.add_paragraph(str(corr))

doc.add_heading("Regression Output", 1)
doc.add_paragraph(str(model.summary()))

doc.save("report.docx")

print("DONE → report.docx created")